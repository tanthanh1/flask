from docx import Document
import re

doc = Document('1.docx')
new_doc = Document()

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # Thêm dấu cách sau chữ hoa nếu thiếu
    text = re.sub(r'([A-Z])\.', r'\1. ', text)

    # Tách theo chữ hoa hoặc số có >=2 chữ số, nhưng số phải đứng sau ký tự không phải số để tránh tách số như 101 thành 1 và 01
    parts = re.split(r'(?=[A-Z]\. )|(?<=\D)(?=\d{2,}\. )', text)

    for part in parts:
        cleaned_part = part.strip()
        if cleaned_part:
            new_doc.add_paragraph(cleaned_part)

new_doc.save('output.docx')
print("Xử lý xong, kết quả lưu trong output.docx")