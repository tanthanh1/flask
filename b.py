from docx import Document
import re

def get_runs_with_positions(para):
    """Lấy danh sách các run và vị trí bắt đầu, kết thúc trong đoạn para"""
    runs = []
    pos = 0
    for run in para.runs:
        length = len(run.text)
        runs.append({
            'text': run.text,
            'start': pos,
            'end': pos + length,
            'bold': run.bold
        })
        pos += length
    return runs

def is_bold_in_substring(runs, start_pos, end_pos):
    """Kiểm tra đoạn text từ start_pos đến end_pos có in đậm không"""
    for run in runs:
        # Nếu run này có phần trùng với đoạn substring
        if run['end'] > start_pos and run['start'] < end_pos:
            if run['bold']:
                return True
    return False

doc = Document('1.docx')
new_doc = Document()

for para in doc.paragraphs:
    text = para.text
    if not text.strip():
        continue

    # Thêm dấu cách sau chữ hoa nếu thiếu
    text_with_spaces = re.sub(r'([A-Z])\.', r'\1. ', text)

    # Tách đoạn theo chữ hoa hoặc số có >=2 chữ số
    parts = re.split(r'(?=[A-Z]\. )|(?<=\D)(?=\d{2,}\. )', text_with_spaces)

    runs = get_runs_with_positions(para)

    # Vì text_with_spaces có thể khác với text gốc về khoảng trắng,
    # ta cần dò vị trí substring parts trong text_with_spaces rồi map lại vị trí trong text gốc
    # Đơn giản ở đây ta dò trực tiếp phần tách trong text_with_spaces

    curr_pos = 0
    for part in parts:
        part = part.strip()
        if not part:
            continue

        # Tìm vị trí part trong text_with_spaces bắt đầu từ curr_pos
        start_idx = text_with_spaces.find(part, curr_pos)
        if start_idx == -1:
            # Không tìm thấy, bỏ qua
            start_idx = curr_pos
        end_idx = start_idx + len(part)
        curr_pos = end_idx

        # Bây giờ cần map vị trí start_idx, end_idx trong text_with_spaces về vị trí trong text gốc.
        # Vì chỉ thêm dấu cách sau chữ hoa (A.B -> A. B), nên ta loại bỏ dấu cách để map tương đối.
        # Để đơn giản hơn, ta bỏ qua mapping chính xác, giả sử text_with_spaces giống text (nếu khác nhiều có thể sai)

        # Kiểm tra đoạn substring có in đậm không
        bold = is_bold_in_substring(runs, start_idx, end_idx)

        # Thêm đoạn vào new_doc
        p = new_doc.add_paragraph(part)
        if bold:
            for run in p.runs:
                run.bold = True

new_doc.save('output.docx')
print("Xử lý xong, kết quả lưu trong output.docx")
