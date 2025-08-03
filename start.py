from docx import Document
import re
import sys

def get_runs_with_positions(para):
    runs = []
    pos = 0
    for run in para.runs:
        length = len(run.text)
        runs.append({
            'text': run.text,
            'start': pos,
            'end': pos + length,
            'bold': run.bold
        })
        pos += length
    return runs

def is_bold_in_substring(runs, start_pos, end_pos):
    for run in runs:
        if run['end'] > start_pos and run['start'] < end_pos:
            if run['bold']:
                return True
    return False

def main():
    if len(sys.argv) < 2:
        print("❌ Cách dùng:")
        print("   python3 start.py input.docx [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "ketqua.docx"

    doc = Document(input_path)
    new_doc = Document()

    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue

        text_with_spaces = re.sub(r'([A-Z])\.', r'\1. ', text)
        parts = re.split(r'(?=[A-Z]\. )|(?<=\D)(?=\d{2,}\. )', text_with_spaces)

        runs = get_runs_with_positions(para)
        curr_pos = 0

        for part in parts:
            part = part.strip()
            if not part:
                continue

            start_idx = text_with_spaces.find(part, curr_pos)
            if start_idx == -1:
                start_idx = curr_pos
            end_idx = start_idx + len(part)
            curr_pos = end_idx

            bold = is_bold_in_substring(runs, start_idx, end_idx)
            p = new_doc.add_paragraph(part)
            if bold:
                for run in p.runs:
                    run.bold = True

    new_doc.save(output_path)
    print(f"✅ Đã xử lý xong! Kết quả lưu trong: {output_path}")

if __name__ == "__main__":
    main()